from fastapi import FastAPI, Request, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from pydantic import BaseModel
from openai import AsyncOpenAI
import os
from dotenv import load_dotenv
from fpdf import FPDF
import io
import json
from docx import Document
from docx.shared import Pt
from bs4 import BeautifulSoup
import aiohttp
import logging
from fastapi.middleware.cors import CORSMiddleware
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import letter

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Загрузка переменных окружения
load_dotenv()

app = FastAPI()

# Настройка CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Монтируем статические файлы
app.mount("/static", StaticFiles(directory="static"), name="static")

# Инициализация шаблонов
templates = Jinja2Templates(directory="templates")


# Модели данных
class TextInput(BaseModel):
    text: str


# Сервис для работы с Mistral AI
class MistralService:
    def __init__(self):
        self.api_key = os.getenv("MISTRAL_API_KEY")
        self.model = "mistral-large-latest"
        self.client = AsyncOpenAI(
            api_key=self.api_key,
            base_url="https://api.mistral.ai/v1"
        )

    async def generate_test(self, text: str):
        try:
            messages = self._create_prompt(text)
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                response_format={"type": "json_object"},
                stream=False,
            )
            return json.loads(response.choices[0].message.content)
        except Exception as e:
            logger.error(f"Ошибка генерации теста: {str(e)}")
            raise HTTPException(status_code=500, detail=str(e))

    def _create_prompt(self, text: str):
        text_length = len(text)

        if text_length < 500:
            num_questions = 3
            num_options = 3
        elif text_length < 1500:
            num_questions = 5
            num_options = 4
        else:
            num_questions = 10
            num_options = 5

        system_prompt = f"""
        Ты - система для создания тестов. Создай тест по следующему тексту:
        - Сгенерируй {num_questions} вопросов
        - В каждом вопросе должно быть {num_options} вариантов ответа
        - Только 1 правильный ответ на вопрос
        - Придумай название теста
        - Верни ответ в JSON формате:

        {{
            "name": "Название теста",
            "questions": [
                {{
                    "question": "Текст вопроса",
                    "options": [
                        {{"answer": "Вариант 1", "correct": true/false}},
                        {{"answer": "Вариант 2", "correct": true/false}}
                    ]
                }}
            ]
        }}
        """

        return [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text}
        ]


# Инициализация сервиса
service = MistralService()


# Маршруты API
@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/api/generate-test")
async def generate_test(text_input: TextInput):
    try:
        test_data = await service.generate_test(text_input.text)

        # Валидация структуры данных
        if not isinstance(test_data, dict):
            raise HTTPException(status_code=500, detail="Некорректный формат теста")

        if 'name' not in test_data:
            test_data['name'] = "Тест без названия"

        if 'questions' not in test_data or not isinstance(test_data['questions'], list):
            raise HTTPException(status_code=500, detail="Некорректная структура вопросов")

        # Валидация каждого вопроса
        for question in test_data['questions']:
            if 'question' not in question or not question['question']:
                question['question'] = "Вопрос без текста"

            if 'options' not in question or not isinstance(question['options'], list):
                question['options'] = []

            for option in question['options']:
                if 'answer' not in option:
                    option['answer'] = "Вариант без текста"
                if 'correct' not in option:
                    option['correct'] = False

        return JSONResponse(content=test_data)

    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error(f"Ошибка генерации теста: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Ошибка генерации теста: {str(e)}")


@app.get("/api/fetch-text")
async def fetch_text_from_url(url: str):
    try:
        # Проверка URL
        if not url.startswith(('http://', 'https://')):
            raise HTTPException(status_code=400, detail="URL должен начинаться с http:// или https://")

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }

        timeout = aiohttp.ClientTimeout(total=10)
        async with aiohttp.ClientSession(headers=headers, timeout=timeout) as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Ошибка загрузки страницы. Код: {response.status}"
                    )

                # Проверка content-type
                content_type = response.headers.get('content-type', '')
                if 'text/html' not in content_type:
                    raise HTTPException(
                        status_code=400,
                        detail="Страница не содержит HTML контент"
                    )

                html = await response.text()
                soup = BeautifulSoup(html, 'html.parser')

                # Удаляем ненужные элементы
                for element in soup(['script', 'style', 'nav', 'footer', 'head', 'iframe', 'img']):
                    element.decompose()

                # Извлекаем текст
                text = soup.get_text('\n')
                text = '\n'.join([line.strip() for line in text.split('\n') if line.strip()])

                if len(text) < 100:
                    raise HTTPException(
                        status_code=400,
                        detail="Не удалось извлечь достаточно текста"
                    )

                return {"text": text}

    except aiohttp.ClientError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Ошибка при запросе к URL: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Неожиданная ошибка: {str(e)}"
        )



@app.post("/api/generate-pdf")
async def generate_pdf(test_data: dict):
    try:
        # Регистрируем шрифт с поддержкой русского языка
        pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))

        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)

        # Настройки
        p.setFont("Arial", 16)
        p.drawString(100, 750, test_data['name'])

        y_position = 700
        p.setFont("Arial", 12)

        for i, question in enumerate(test_data['questions'], 1):
            p.drawString(100, y_position, f"{i}. {question['question']}")
            y_position -= 20

            for j, option in enumerate(question['options'], 1):
                p.drawString(120, y_position, f"{j}. {option['answer']}")
                y_position -= 15

            y_position -= 10

            if y_position < 50:
                p.showPage()
                y_position = 750
                p.setFont("Arial", 12)

        p.save()
        buffer.seek(0)

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=test.pdf"}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/generate-word")
async def generate_word(test_data: dict):
    try:
        # Создаем документ Word
        doc = Document()

        # Добавляем название теста
        title = doc.add_heading(test_data.get('name', 'Тест без названия'), level=1)
        title.alignment = 1  # Центральное выравнивание
        doc.add_paragraph()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # Добавляем вопросы
        for i, question in enumerate(test_data.get('questions', []), 1):
            # Вопрос
            q_paragraph = doc.add_paragraph()
            q_paragraph.add_run(f"{i}. {question.get('question', 'Вопрос без текста')}").bold = True

            # Варианты ответов
            for j, option in enumerate(question.get('options', []), 1):
                doc.add_paragraph(f"   {j}. {option.get('answer', 'Вариант без текста')}")

            doc.add_paragraph()  # Пустая строка между вопросами

        # Возвращаем Word как поток
        word_buffer = io.BytesIO()
        doc.save(word_buffer)
        word_buffer.seek(0)

        return StreamingResponse(
            word_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=test.docx"}
        )
    except Exception as e:
        logger.error(f"Ошибка генерации Word: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="127.0.0.1", port=8000, reload=True)